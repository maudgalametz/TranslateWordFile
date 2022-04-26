#%%
from cmath import nan
import docx
import pandas as pd

document = docx.Document('Test.docx') 
exc = pd.read_excel('Test.xlsx')


# %%
import re
from docx import Document

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text, 1)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


#%%
for texte_ in range(len(exc["Anglais"])):
    if (not(pd.isnull(exc["Français"][texte_]))) & (not(pd.isnull(exc["Français"][texte_]))):
        text_old = re.compile(exc["Français"][texte_])
        text_new = exc["Anglais"][texte_]
        docx_replace_regex(document, text_old , text_new)


document.save('FileTranslated.docx')
# %%
